""""""
import docx
import docx2txt

import logging

logger = logging.getLogger(__name__)

class WordParser:
    """
    
    ---
    Attributes
    ----------


    Methods
    -------

    
    """
    
    name = "Word Parser"
    version = "0.0.1"

    def __init__(self):
        pass

    def __str__(self):
        return "{} {}".format(self.name, self.version)
        
    def get_extensions(self):
        """
        
        Parameters
        ----------

        Returns
        -------
        None
        """
        return [".docx"]
        
    def get_functions(self):
        """
        
        Parameters
        ----------

        Returns
        -------
        None
        """
        return {
            "metadata": self.get_metadata,
            "analyse": self.analyze,
            "contents": self.get_contents
        }

    def get_metadata(self, filepath: str) -> dict:
        """This will return the doc info infomation from the 
        Named file.
        
        Parameters
        ----------

        Returns
        -------
        None
        """

        data = {}
        doc = docx.Document(filepath)
        
        # get the core properties from the file...
        # https://python-docx.readthedocs.io/en/latest/api/document.html#coreproperties-objects
        cp = doc.core_properties
        
        data['author'] = cp.author
        data['category'] = cp.category
        data['comments'] = cp.comments
        data['content_status'] = cp.content_status
        data['created'] = cp.created
        data['identifier'] = cp.identifier
        data['keywords'] = cp.keywords
        data['language'] = cp.language
        data['last_modified_by'] = cp.last_modified_by
        data['last_printed'] = cp.last_printed
        data['modified'] = cp.modified
        data['revision'] = cp.revision
        data['subject'] = cp.subject
        data['title'] = cp.title
        data['version'] = cp.version
        
        return data

    def analyze(self, filepath: str) -> dict:
        """
        
        Parameters
        ----------

        Returns
        -------
        None
        """
        data = dict
        
        return data

    def get_contents(self, filepath: str) -> list:
        """This will return the paragroah objects in a word document
        
        Parameters
        ----------

        Returns
        -------
        None
        """
        text = docx2txt.process(filepath)
        return {
            "type": "strings",
            "strings": text
        }

    def get_contentsD(self, filepath: str) -> list:
        """This will return the paragroah objects in a word document
        
        Parameters
        ----------

        Returns
        -------
        None
        """
        data = []
        doc = docx.Document(filepath)

        for p in doc.paragraphs:
            data.append(p.text)

        return data

